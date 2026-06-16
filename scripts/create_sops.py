"""
create_sops.py
--------------
Creates 4 bank backend SOP Word documents in the data/ folder.

Usage:
    python scripts/create_sops.py
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.makedirs("data", exist_ok=True)

def make_doc(): return Document()
def h2(doc, t): doc.add_heading(t, level=2)
def h3(doc, t): doc.add_heading(t, level=3)
def para(doc, t): doc.add_paragraph(t)
def sp(doc): doc.add_paragraph()
def bullet(doc, items):
    for i in items: doc.add_paragraph(i, style="List Bullet")
def numbered(doc, items):
    for i in items: doc.add_paragraph(i, style="List Number")

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        c.paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row): t.rows[ri].cells[ci].text = v
    sp(doc)

def title_block(doc, title, subtitle, meta):
    x = doc.add_heading(title, level=0); x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph(subtitle); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].bold = True; s.runs[0].font.size = Pt(13); sp(doc)
    tbl(doc, ["Field","Details"], meta)


# ── SOP 1: Loan Origination & Underwriting ─────────────────────────────
def sop1():
    doc = make_doc()
    title_block(doc, "Standard Operating Procedure", "Loan Origination & Underwriting Process", [
        ("Document ID","SOP-BANK-LO-012"),("Version","2.4"),("Effective Date","01-June-2026"),
        ("Review Date","01-June-2027"),("Owner","Chief Credit Officer"),
        ("Department","Retail & Commercial Lending"),("Frequency","Per loan application received"),
        ("Applies To","Consumer Loans | Auto Loans | Business Loans | Mortgages")])

    h2(doc,"1. Purpose")
    para(doc,"This SOP defines the end-to-end process for receiving, evaluating, approving, and disbursing loan applications. It ensures consistent credit decisions, regulatory compliance, and operational efficiency across all loan product types.")

    h2(doc,"2. Scope")
    bullet(doc,["Branch walk-in applications","Digital / online portal submissions","Dealer-submitted applications (Indirect Auto)","Broker-submitted commercial applications"])

    h2(doc,"3. Roles & Responsibilities")
    tbl(doc,["Role","Responsibility","Contact"],[
        ("Loan Officer (LO)","Collects application, verifies documents, submits to underwriting","lo-team@bank.com"),
        ("Underwriter (L1)","Reviews creditworthiness, runs models, issues preliminary decision","underwriting@bank.com"),
        ("Senior Underwriter (L2)","Handles exceptions, large loans, and complex cases","sr.underwriting@bank.com"),
        ("Credit Committee","Approves loans above $500,000 or with exception requests","creditcommittee@bank.com"),
        ("Loan Processor","Prepares closing documents, schedules disbursement","loanops@bank.com"),
        ("Compliance Officer","Reviews for ECOA, HMDA, fair lending compliance","compliance@bank.com"),
        ("Chief Credit Officer","Final escalation authority; approves policy exceptions","cco@bank.com")])

    h2(doc,"4. Step-by-Step Process")
    h3(doc,"Step 1 – Application Receipt & Initial Review")
    numbered(doc,["Loan Officer (LO) collects completed application form with all required documents.",
        "Required documents: Government ID, last 2 pay stubs, last 2 years tax returns (business loans), bank statements (3 months), collateral details (if secured).",
        "LO logs application in Loan Origination System (LOS) — Encompass or nCino — within 2 hours of receipt.",
        "System auto-generates Application ID (e.g., APP-2026-08821).",
        "Send acknowledgement to applicant within 24 hours with Application ID and expected decision timeline.",
        "If documents are incomplete: send Document Deficiency Notice within 1 business day. Application goes on hold."])

    h3(doc,"Step 2 – Credit Pull & Preliminary Screening")
    numbered(doc,["LO pulls credit report from all three bureaus (Equifax, Experian, TransUnion) via LOS.",
        "System auto-screens for: minimum FICO score, debt-to-income ratio (DTI), existing delinquencies.",
        "Auto-decline triggers: FICO < 580, DTI > 55%, active bankruptcy, fraud alert on file.",
        "If auto-decline: generate Adverse Action Notice (AAN) within 30 days per Reg B. Send to applicant.",
        "If passes screening: assign to underwriting queue. SLA: underwriter picks up within 4 business hours."])

    h3(doc,"Step 3 – Underwriting & Credit Decision")
    numbered(doc,["Underwriter reviews full file: credit report, income verification, collateral appraisal (if applicable).",
        "Run debt service coverage ratio (DSCR) for business loans — minimum 1.25x required.",
        "Check LTV: max 80% for real estate, 90% for auto loans with PMI.",
        "Document findings in underwriting memo in LOS.",
        "Issue one of three decisions: APPROVE / CONDITIONAL APPROVE / DECLINE.",
        "Decision SLA: 2 business days for consumer loans; 5 business days for business loans."])

    h3(doc,"Step 4 – Approval & Loan Committee Review")
    numbered(doc,["Loans above $500,000 must go to Credit Committee — submit file at least 2 business days before meeting.",
        "Credit Committee meets every Tuesday and Thursday at 10:00 AM.",
        "Committee decision documented in minutes and uploaded to LOS within 24 hours.",
        "Policy exceptions require Credit Committee + CCO sign-off."])

    h3(doc,"Step 5 – Loan Closing & Disbursement")
    numbered(doc,["Loan Processor prepares closing package: promissory note, security agreement, disclosure forms.",
        "Schedule closing appointment with applicant — in branch or via e-sign platform.",
        "All closing documents reviewed by Compliance for TRID, RESPA, or Reg Z compliance.",
        "Disbursement authorized only after: all documents signed, conditions cleared, compliance sign-off received.",
        "Funds disbursed via ACH (consumer) or wire transfer (business) within 1 business day of closing.",
        "Loan boarded in core banking system (FiServ/Jack Henry) by Loan Processor same day."])

    h2(doc,"5. What To Do If Stuck")
    tbl(doc,["Situation","Action"],[
        ("Applicant disputes auto-decline","Review credit report for errors; advise applicant to dispute with bureau; do not override without Sr. Underwriter approval"),
        ("Collateral appraisal lower than expected","Order second appraisal; if still low, reduce loan amount or require additional collateral"),
        ("Applicant cannot provide tax returns","Accept CPA-prepared P&L for <2yr businesses; escalate exception to Sr. Underwriter"),
        ("Disbursement delayed by IT issue","Contact loanops@bank.com; notify applicant; document delay in LOS"),
        ("Fraud indicators found during review","Immediately suspend application; notify fraud@bank.com; do not inform applicant"),
        ("Application exceeds $500K threshold","Route to Credit Committee; do not approve at underwriter level")])

    h2(doc,"6. Key Contacts")
    tbl(doc,["Role","Email","Phone"],[
        ("Loan Operations","loanops@bank.com","+1-800-555-0201"),
        ("Underwriting","underwriting@bank.com","+1-800-555-0202"),
        ("Compliance","compliance@bank.com","+1-800-555-0199"),
        ("Fraud Team","fraud@bank.com","+1-800-555-0188"),
        ("Chief Credit Officer","cco@bank.com","+1-800-555-0210")])

    h2(doc,"7. Change Log")
    tbl(doc,["Version","Date","Author","Changes"],[
        ("1.0","Jan 2023","CCO Office","Initial version"),
        ("2.0","Aug 2024","Compliance","Added TRID and RESPA requirements; updated DTI thresholds"),
        ("2.4","Jun 2026","Underwriting","Added nCino LOS steps; updated Credit Committee thresholds")])

    doc.save(os.path.join("data","SOP_Loan_Origination_Underwriting.docx"))
    print("✅ Saved: SOP_Loan_Origination_Underwriting.docx")


# ── SOP 2: Wire Transfer & ACH Processing ─────────────────────────────
def sop2():
    doc = make_doc()
    title_block(doc,"Standard Operating Procedure","Wire Transfer & ACH Payment Processing",[
        ("Document ID","SOP-BANK-PAY-009"),("Version","3.1"),("Effective Date","01-June-2026"),
        ("Review Date","01-June-2027"),("Owner","Head of Payment Operations"),
        ("Department","Payment Operations & Treasury"),("Frequency","Daily — multiple times per business day"),
        ("Applies To","Domestic Wire | International Wire | ACH Credit | ACH Debit")])

    h2(doc,"1. Purpose")
    para(doc,"This SOP governs the initiation, processing, verification, and settlement of wire transfers and ACH transactions. It ensures accuracy, regulatory compliance, OFAC screening, and fraud prevention for all payment types.")

    h2(doc,"2. Payment Cut-Off Times")
    tbl(doc,["Payment Type","Cut-Off Time","Settlement","System"],[
        ("Domestic Wire (Fedwire)","4:30 PM EST","Same day","FedLine Advantage"),
        ("International Wire (SWIFT)","2:00 PM EST","1-2 business days","SWIFT / Correspondent"),
        ("ACH Credit (same-day)","2:45 PM EST","Same day","FedACH / NACHA"),
        ("ACH Credit (next-day)","11:59 PM EST","Next business day","FedACH / NACHA"),
        ("ACH Debit","11:59 PM EST","Next business day","FedACH / NACHA"),
        ("Internal Transfer","5:00 PM EST","Immediate","Core Banking System")])

    h2(doc,"3. Roles & Responsibilities")
    tbl(doc,["Role","Responsibility","Contact"],[
        ("Payment Initiator","Inputs payment details; submits for dual approval","payments@bank.com"),
        ("Payment Approver","Second-level review; authorizes release","pay.approval@bank.com"),
        ("OFAC Screening Team","Reviews flagged transactions; clears or escalates","ofac@bank.com"),
        ("Fraud Operations","Investigates suspicious patterns; freezes fraudulent payments","fraud@bank.com"),
        ("Treasury / Settlements","Monitors settlement; manages intraday liquidity","treasury@bank.com"),
        ("Head of Payment Ops","Escalation authority; approves out-of-policy payments","h.payops@bank.com")])

    h2(doc,"4. Step-by-Step Process")
    h3(doc,"Step 1 – Payment Initiation")
    numbered(doc,["Log into payment system (FIS Global PAYplus).",
        "Enter: beneficiary name, account number, routing/SWIFT number, amount, currency, payment date, purpose code.",
        "For wire transfers over $10,000: attach supporting documentation (invoice, contract, authorization letter).",
        "Submit payment for dual approval — do NOT self-approve under any circumstance.",
        "System auto-generates Payment Reference ID (e.g., PAY-2026-044521)."])

    h3(doc,"Step 2 – OFAC & Sanctions Screening")
    numbered(doc,["All payments automatically screened against OFAC SDN list, EU and UN sanctions lists upon submission.",
        "If CLEARED by system: proceed to Step 3.",
        "If FLAGGED: system places payment on hold automatically.",
        "OFAC Screening Team reviews flagged payment within 1 hour.",
        "If false positive: OFAC team clears hold with documented reason.",
        "If true match: block payment immediately; notify compliance@bank.com and Legal within 15 minutes; file OFAC report with FinCEN.",
        "Do NOT release a true OFAC match without Legal sign-off."])

    h3(doc,"Step 3 – Dual Approval")
    numbered(doc,["Payment Approver independently verifies: beneficiary details, amount, account numbers, supporting documents.",
        "Approver must NOT approve if: details cannot be verified, payment appears unusual, or urgency claimed without documentation.",
        "Payments over $1,000,000 require VP-level approver."])

    h3(doc,"Step 4 – Settlement & Confirmation")
    numbered(doc,["Domestic wires sent via Fedwire; confirmation received within minutes.",
        "International wires sent via SWIFT MT103; confirmation via MT199 from correspondent bank.",
        "ACH submitted to FedACH as NACHA file; settlement confirmed next morning.",
        "Send payment confirmation to initiator and account holder via system notification.",
        "Log all settlement confirmations in Payment Register by end of day."])

    h3(doc,"Step 5 – Return & Rejection Handling")
    numbered(doc,["ACH returns received by 8:00 AM next business day — review in FedACH portal.",
        "Common return codes: R01 (insufficient funds), R02 (account closed), R03 (no account), R10 (not authorized).",
        "For R10 returns: freeze originating account; notify Fraud Ops immediately.",
        "Wire rejections via SWIFT MT199: identify reason, correct, resubmit same day if cut-off not passed.",
        "Notify account holder of any return within 2 business hours.",
        "Escalate patterns (3+ returns from same originator) to Fraud Ops."])

    h2(doc,"5. What To Do If Stuck")
    tbl(doc,["Situation","Action"],[
        ("Payment stuck in OFAC review > 2 hours","Escalate to compliance@bank.com and h.payops@bank.com"),
        ("Missed cut-off for urgent wire","Contact treasury@bank.com for emergency Fedwire; requires VP approval"),
        ("Beneficiary bank returns wire — unknown reason","Contact correspondent@bank.com; request SWIFT trace via MT199"),
        ("System (FIS) unavailable during processing","Switch to manual backup procedures; notify treasury@bank.com; call IT on-call via PagerDuty"),
        ("Customer claims payment not received","Pull Fedwire/SWIFT confirmation; if confirmed sent, advise customer to contact their bank")])

    h2(doc,"6. Key Contacts")
    tbl(doc,["Role","Email","Phone"],[
        ("Payment Operations","payments@bank.com","+1-800-555-0301"),
        ("OFAC Team","ofac@bank.com","+1-800-555-0302"),
        ("Fraud Operations","fraud@bank.com","+1-800-555-0188"),
        ("Treasury","treasury@bank.com","+1-800-555-0303"),
        ("Head of Payment Ops","h.payops@bank.com","+1-800-555-0310")])

    h2(doc,"7. Change Log")
    tbl(doc,["Version","Date","Author","Changes"],[
        ("1.0","Mar 2021","Payment Ops","Initial version"),
        ("2.0","Jan 2024","Compliance","Added same-day ACH; updated OFAC escalation procedures"),
        ("3.1","Jun 2026","Payment Ops","Updated cut-off times; added international wire details")])

    doc.save(os.path.join("data","SOP_Wire_ACH_Payment_Processing.docx"))
    print("✅ Saved: SOP_Wire_ACH_Payment_Processing.docx")


# ── SOP 3: KYC & AML Compliance ────────────────────────────────────────
def sop3():
    doc = make_doc()
    title_block(doc,"Standard Operating Procedure","Know Your Customer (KYC) & Anti-Money Laundering (AML) Compliance",[
        ("Document ID","SOP-BANK-AML-003"),("Version","4.0"),("Effective Date","01-June-2026"),
        ("Review Date","01-June-2027"),("Owner","Chief Compliance Officer (CCO)"),
        ("Department","Compliance & Risk"),("Frequency","At onboarding + periodic reviews + triggered by alerts"),
        ("Applies To","All customer types: Retail | Business | High-Net-Worth | Foreign Nationals")])

    h2(doc,"1. Purpose")
    para(doc,"This SOP defines the bank's KYC and AML procedures including customer identification, due diligence, enhanced due diligence, transaction monitoring, suspicious activity reporting, and customer risk rating. It ensures compliance with the Bank Secrecy Act (BSA), FinCEN requirements, USA PATRIOT Act, and FATF guidelines.")

    h2(doc,"2. Customer Risk Tiers")
    tbl(doc,["Risk Tier","Customer Profile","KYC Review Frequency","EDD Required?"],[
        ("Tier 1 – Low Risk","Domestic retail customers, salaried employees, standard accounts","Every 3 years","No"),
        ("Tier 2 – Medium Risk","Small businesses, sole proprietors, non-US residents","Every 2 years","Conditional"),
        ("Tier 3 – High Risk","Cash-intensive businesses, foreign nationals, PEPs, high transaction volume","Every 1 year","Yes"),
        ("Tier 4 – Prohibited","OFAC-sanctioned entities, known fraud/AML convictions","N/A – Decline","N/A – Decline")])

    h2(doc,"3. Roles & Responsibilities")
    tbl(doc,["Role","Responsibility","Contact"],[
        ("KYC Analyst","Collects and verifies customer identity documents at onboarding","kyc@bank.com"),
        ("AML Investigator","Reviews transaction monitoring alerts; investigates suspicious activity","aml@bank.com"),
        ("BSA Officer","Oversees SAR/CTR filing; primary FinCEN liaison","bsa@bank.com"),
        ("Relationship Manager","Collects business documentation; assists with EDD interviews","rm@bank.com"),
        ("Chief Compliance Officer","Final sign-off on SAR decisions; escalation authority","cco@bank.com"),
        ("Legal","Reviews complex SAR cases; manages government requests","legal@bank.com")])

    h2(doc,"4. Step-by-Step Process")
    h3(doc,"Step 1 – Customer Identification Program (CIP)")
    numbered(doc,["Collect: Full legal name, Date of birth (individuals), Address, Tax ID (SSN/EIN/ITIN).",
        "Verify identity using: Government-issued photo ID, SSN verification via Equifax/LexisNexis, Business registration documents (for entities).",
        "For business accounts additionally collect: Articles of Incorporation, Beneficial Ownership form (all owners with >25% stake), Operating Agreement.",
        "Screen all names against OFAC SDN list, PEP database, and internal watchlist.",
        "Document all CIP information in KYC system (Actimize) within 1 business day.",
        "If identity cannot be verified within 5 business days: do not open account; notify customer."])

    h3(doc,"Step 2 – Risk Rating & Due Diligence Level")
    numbered(doc,["Assign risk tier (1-4) based on: customer type, geography, business type, transaction profile.",
        "Tier 1-2: Standard Due Diligence (SDD) — CIP documents sufficient.",
        "Tier 3: Enhanced Due Diligence (EDD) required — see Step 3.",
        "Tier 4: Decline customer; log reason in KYC system; notify CCO."])

    h3(doc,"Step 3 – Enhanced Due Diligence (EDD)")
    numbered(doc,["EDD triggers: PEP status, high-risk country, cash-intensive business, unusual account purpose.",
        "EDD steps: In-person or video interview, source of wealth documentation, 3 months of bank statements, Senior VP approval before account opening.",
        "EDD file reviewed by BSA Officer before account is approved.",
        "EDD customers reviewed annually — schedule review 30 days before anniversary."])

    h3(doc,"Step 4 – Transaction Monitoring")
    numbered(doc,["Automated monitoring (Actimize) generates alerts for: structuring patterns (transactions just below $10,000), sudden spike in transaction volume, wire transfers to high-risk jurisdictions.",
        "AML Investigator reviews alerts within 5 business days.",
        "Alert outcomes: CLEARED (document reason), ESCALATED TO SAR REVIEW, or ACCOUNT RESTRICTION.",
        "Accounts with 3+ cleared alerts in 90 days are auto-escalated for risk re-rating."])

    h3(doc,"Step 5 – Currency Transaction Report (CTR) Filing")
    numbered(doc,["CTR required for all cash transactions exceeding $10,000 in a single business day.",
        "CTR must be filed with FinCEN within 15 calendar days of transaction.",
        "BSA Officer reviews and submits via BSA E-Filing System.",
        "Do NOT tip off the customer that a CTR is being filed."])

    h3(doc,"Step 6 – Suspicious Activity Report (SAR) Filing")
    numbered(doc,["SAR required when suspicious activity involves $5,000 or more.",
        "SAR decision made jointly by AML Investigator and BSA Officer.",
        "SAR must be filed within 30 days of initial detection (60 days if no suspect identified).",
        "File via FinCEN BSA E-Filing System — retain copy for 5 years.",
        "CRITICAL: Do NOT inform the customer or any third party that a SAR has been filed (tipping off is a federal crime).",
        "Notify CCO and Legal of all SAR filings on same day as submission."])

    h2(doc,"5. What To Do If Stuck")
    tbl(doc,["Situation","Action"],[
        ("Customer refuses Beneficial Ownership form","Cannot open business account without it; inform customer it is legally required under FinCEN rules"),
        ("OFAC match on existing customer","Freeze account immediately; notify ofac@bank.com, compliance@bank.com, and Legal within 1 hour"),
        ("Unsure if activity warrants a SAR","When in doubt, file the SAR — consult BSA Officer; law protects good-faith filers"),
        ("Law enforcement arrives with subpoena","Direct to Legal immediately; do not provide records without Legal review"),
        ("Customer shows structuring pattern","Do not alert customer; file SAR; consider account restriction pending investigation")])

    h2(doc,"6. Key Contacts")
    tbl(doc,["Role","Email","Phone"],[
        ("KYC Team","kyc@bank.com","+1-800-555-0401"),
        ("AML Team","aml@bank.com","+1-800-555-0402"),
        ("BSA Officer","bsa@bank.com","+1-800-555-0403"),
        ("Chief Compliance Officer","cco@bank.com","+1-800-555-0199"),
        ("Legal","legal@bank.com","+1-800-555-0450"),
        ("FinCEN","Via BSA E-Filing","1-800-767-2825")])

    h2(doc,"7. Change Log")
    tbl(doc,["Version","Date","Author","Changes"],[
        ("1.0","Jun 2019","BSA Office","Initial version"),
        ("2.0","Jan 2021","Compliance","Added Beneficial Ownership rules per FinCEN CDD Rule"),
        ("3.0","Mar 2024","CCO Office","Updated EDD requirements; added Actimize monitoring rules"),
        ("4.0","Jun 2026","BSA Officer","Updated SAR thresholds; added PEP screening; revised risk tier definitions")])

    doc.save(os.path.join("data","SOP_KYC_AML_Compliance.docx"))
    print("✅ Saved: SOP_KYC_AML_Compliance.docx")


# ── SOP 4: Core Banking EOD Processing ────────────────────────────────
def sop4():
    doc = make_doc()
    title_block(doc,"Standard Operating Procedure","Core Banking System – End-of-Day (EOD) Processing",[
        ("Document ID","SOP-BANK-EOD-001"),("Version","5.2"),("Effective Date","01-June-2026"),
        ("Review Date","01-June-2027"),("Owner","VP of IT Operations"),
        ("Department","IT Operations & Core Banking"),("Frequency","Daily — every business day after 5:00 PM EST"),
        ("Applies To","FiServ Precision | Jack Henry Symitar | All Ancillary Systems")])

    h2(doc,"1. Purpose")
    para(doc,"This SOP defines the daily end-of-day batch processing sequence for the core banking system. It ensures all transactions are posted, accounts are balanced, interest is accrued, reports are generated, and the system is prepared for the next business day — with clear escalation steps if any job fails.")

    h2(doc,"2. EOD Job Sequence")
    tbl(doc,["Seq","Job Name","Start Time","Duration","System","Owner"],[
        ("1","Transaction cutoff — freeze new postings","5:00 PM","5 min","Core Banking","Batch Operator"),
        ("2","Teller & branch balancing","5:05 PM","15 min","Core Banking","Branch Ops"),
        ("3","ACH return file processing","5:20 PM","20 min","FedACH Interface","Payment Ops"),
        ("4","Loan payment posting batch","5:40 PM","30 min","Loan Servicing","Batch Operator"),
        ("5","Interest accrual calculation","6:10 PM","25 min","Core Banking","Batch Operator"),
        ("6","Fee posting (NSF, maintenance, late fees)","6:35 PM","15 min","Core Banking","Batch Operator"),
        ("7","GL posting and trial balance","6:50 PM","20 min","GL System","GL Recon Team"),
        ("8","GL balancing check — must be ZERO","7:10 PM","10 min","GL System","GL Recon Team"),
        ("9","Card authorization batch settlement","7:20 PM","30 min","Card Processing","Cards Team"),
        ("10","Regulatory report extract","7:50 PM","20 min","Reporting System","Compliance"),
        ("11","Statement generation","8:10 PM","45 min","Core Banking","Batch Operator"),
        ("12","Ancillary system interface sync","8:55 PM","20 min","All interfaces","Core Banking Admin"),
        ("13","Full system backup","9:15 PM","60 min","Backup Systems","IT Operations"),
        ("14","System date roll to next business day","10:15 PM","5 min","Core Banking","Batch Operator"),
        ("15","EOD completion report generated","10:20 PM","5 min","Reporting System","Batch Operator")])

    h2(doc,"3. Roles & Responsibilities")
    tbl(doc,["Role","Responsibility","Contact"],[
        ("EOD Batch Operator","Initiates and monitors EOD batch job sequence","eod.ops@bank.com"),
        ("Core Banking Admin","Resolves job failures; manages system parameters","corebanking@bank.com"),
        ("GL Reconciliation Team","Balances general ledger; investigates out-of-balance conditions","gl.recon@bank.com"),
        ("IT Operations Manager","Escalation point for P1 batch failures; coordinates recovery","it.manager@bank.com"),
        ("VP of IT Operations","Final escalation; approves emergency rollback or recovery","vp.itops@bank.com"),
        ("Business Continuity Team","Activates DR plan if primary system fails","bcm@bank.com")])

    h2(doc,"4. Step-by-Step Process")
    h3(doc,"Step 1 – Pre-EOD Preparation (4:00 PM – 5:00 PM)")
    numbered(doc,["Confirm all branches submitted teller balancing by 4:45 PM.",
        "Check for pending manual journal entries — must be approved and posted before 5:00 PM cutoff.",
        "Confirm ACH files from FedACH have been received and staged.",
        "Verify no open incidents in monitoring dashboard (Dynatrace/Nagios).",
        "Send EOD start notification to all department heads at 5:00 PM."])

    h3(doc,"Step 2 – Run EOD Batch Sequence (5:00 PM – 10:20 PM)")
    numbered(doc,["Log into core banking operations console using EOD service account.",
        "Execute jobs in the sequence listed in Section 2 — do NOT skip or reorder jobs.",
        "After each job: verify status is SUCCESS in job scheduler (Control-M).",
        "If any job shows WARNING: note it, continue sequence, investigate after completion.",
        "If any job shows FAILURE: STOP sequence immediately — escalate per Step 3."])

    h3(doc,"Step 3 – Job Failure Escalation Procedure")
    numbered(doc,["Immediately stop EOD batch sequence.",
        "Capture full error message and job logs from Control-M.",
        "Notify Core Banking Admin within 5 minutes of failure.",
        "Core Banking Admin has 20 minutes to diagnose and resolve.",
        "If not resolved in 20 minutes: escalate to IT Operations Manager.",
        "IT Operations Manager decides: fix and re-run, skip job (if non-critical), or activate recovery procedure.",
        "VP of IT Operations must be notified for any failure that delays EOD past midnight."])

    h3(doc,"Step 4 – GL Balancing (Job 8)")
    numbered(doc,["GL Reconciliation Team pulls trial balance report after Job 7 completes.",
        "Trial balance must net to ZERO — any out-of-balance is a P1 incident.",
        "If out-of-balance: halt remaining EOD jobs; notify IT Operations Manager and CFO immediately.",
        "GL must be in balance before system date roll (Job 14) — no exceptions.",
        "If not resolved by 11:00 PM: CFO and VP of IT Ops jointly approve proceeding with suspense entry."])

    h3(doc,"Step 5 – EOD Completion & Verification")
    numbered(doc,["After Job 15: Batch Operator reviews completion report for anomalies.",
        "Verify: account count matches previous day, total deposit balance within expected range, all interfaces show SUCCESS.",
        "Send EOD Completion Email to IT Ops Manager, CFO Office, and department heads by 10:30 PM.",
        "File EOD Log in SharePoint: IT > EOD Logs > [Date]. Retain for 3 years."])

    h3(doc,"Step 6 – Next Business Day Preparation")
    numbered(doc,["Confirm system date has rolled to next business day.",
        "Run morning readiness check at 7:00 AM: verify all interfaces online, online banking portal accessible, ATM network active.",
        "Send Morning Readiness Report to IT Ops Manager and Branch Operations by 7:30 AM.",
        "Branch tellers should not begin transactions until Morning Readiness confirmation received."])

    h2(doc,"5. What To Do If Stuck")
    tbl(doc,["Situation","Action"],[
        ("Job failure with unknown error","Capture logs; contact corebanking@bank.com; do not retry without guidance"),
        ("EOD not completed by midnight","Notify VP of IT Ops and CFO; activate Business Continuity Plan"),
        ("GL out-of-balance, cause unknown","Engage GL Recon Team + Core Banking Admin jointly; escalate to CFO if not resolved by 11 PM"),
        ("Backup job fails","Do not roll system date until backup succeeds or VP of IT Ops approves exception"),
        ("Core banking system unresponsive","Activate DR failover checklist; notify all department heads; do not attempt manual transactions")])

    h2(doc,"6. Key Contacts")
    tbl(doc,["Role","Email","Phone"],[
        ("EOD Batch Operations","eod.ops@bank.com","+1-800-555-0501"),
        ("Core Banking Admin","corebanking@bank.com","+1-800-555-0502"),
        ("GL Reconciliation","gl.recon@bank.com","+1-800-555-0503"),
        ("IT Operations Manager","it.manager@bank.com","+1-800-555-0510"),
        ("VP of IT Operations","vp.itops@bank.com","+1-800-555-0515"),
        ("On-Call IT","PagerDuty","Via PagerDuty app")])

    h2(doc,"7. Change Log")
    tbl(doc,["Version","Date","Author","Changes"],[
        ("1.0","Apr 2018","IT Ops","Initial version"),
        ("3.0","Feb 2022","Core Banking","Migrated to FiServ Precision"),
        ("5.2","Jun 2026","VP IT Ops","Added DR failover procedures; updated GL balancing P1 definition")])

    doc.save(os.path.join("data","SOP_Core_Banking_EOD_Processing.docx"))
    print("✅ Saved: SOP_Core_Banking_EOD_Processing.docx")


if __name__ == "__main__":
    sop1(); sop2(); sop3(); sop4()
    print("\n🎉 All 4 SOPs created in data/ folder!")
